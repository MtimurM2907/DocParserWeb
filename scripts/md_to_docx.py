#!/usr/bin/env python3
"""Convert DocParseLab markdown explanatory note to .docx."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 14)]:
        if name in doc.styles:
            h = doc.styles[name]
            h.font.name = "Times New Roman"
            h.font.size = Pt(size)
            h.font.bold = True
            h.paragraph_format.first_line_indent = Cm(0)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None, indent: bool = True) -> None:
    p = doc.add_paragraph(style=style)
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(12)
        else:
            run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _format_cell_text(cell, *, bold: bool = False, size: int = 12) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ""
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 0, 0)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def add_markdown_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        _format_cell_text(cell, bold=True)
        _set_cell_shading(cell, "D9E2F3")

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            val = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            _format_cell_text(cell, bold=False)

    doc.add_paragraph()


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_default_font(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            add_formatted_paragraph(doc, stripped[2:], style="Heading 1", indent=False)
            i += 1
            continue

        if stripped.startswith("## "):
            add_formatted_paragraph(doc, stripped[3:], style="Heading 1", indent=False)
            i += 1
            continue

        if stripped.startswith("### "):
            add_formatted_paragraph(doc, stripped[4:], style="Heading 2", indent=False)
            i += 1
            continue

        if stripped.startswith("#### "):
            add_formatted_paragraph(doc, stripped[5:], style="Heading 3", indent=False)
            i += 1
            continue

        if stripped.startswith("> "):
            add_formatted_paragraph(doc, stripped[2:], indent=True)
            p = doc.paragraphs[-1]
            p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            if i > 0:
                prev = lines[i - 1].strip()
                if re.match(r"^\*\*Таблица\s+\d+", prev):
                    cap_text = prev[2:-2] if prev.startswith("**") and prev.endswith("**") else prev
                    add_table_caption(doc, cap_text)

            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, headers, rows)
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                elif part:
                    r = p.add_run(part)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            text = stripped[2:]
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                elif part:
                    r = p.add_run(part)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(11)
            continue

        if not stripped:
            i += 1
            continue

        add_formatted_paragraph(doc, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "DocParseLab_пояснительная_записка.md"
    out = root / "DocParseLab_пояснительная_записка.docx"
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
        if not out.is_absolute():
            out = root / out
    if len(sys.argv) > 2:
        md_arg = Path(sys.argv[2])
        md = md_arg if md_arg.is_absolute() else root / md_arg
    convert(md, out)
